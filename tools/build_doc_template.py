#!/usr/bin/env python3
"""Generate the Digi Word template with brand styles.

Regenerate with:  python3 tools/build_doc_template.py
The output is committed so consumers do not need python-docx.

Regeneration is not byte-deterministic: python-docx embeds timestamps and
revision IDs, so re-running this produces a different file even when no style
changed. Review a regenerated template by reading its styles back (see
tests/test_skills.py), never by byte-comparing it against the committed copy.
Only commit a regenerated binary when a style actually changed.
"""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "assets" / "digi-doc-template.docx"

NAVY = RGBColor(0x1B, 0x49, 0x65)
DARK_GRAY = RGBColor(0x3F, 0x42, 0x45)
TEAL = RGBColor(0x1F, 0x7F, 0xA5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Source Sans Pro"


def add_style(doc, name, size_pt, color, bold=False):
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = FONT
    style.font.size = Pt(size_pt)
    style.font.color.rgb = color
    style.font.bold = bold
    return style


def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK_GRAY

    add_style(doc, "Digi Title", 28, NAVY, bold=True)
    add_style(doc, "Digi Heading 1", 18, NAVY, bold=True)
    add_style(doc, "Digi Heading 2", 14, TEAL, bold=True)
    add_style(doc, "Digi Body", 11, DARK_GRAY)
    add_style(doc, "Digi Table Header", 11, WHITE, bold=True)

    doc.add_paragraph("Digi Document Title", style="Digi Title")
    doc.add_paragraph("Section heading", style="Digi Heading 1")
    doc.add_paragraph("Subsection heading", style="Digi Heading 2")
    doc.add_paragraph(
        "Body text in Source Sans Pro. Replace this content; keep the styles.",
        style="Digi Body",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
